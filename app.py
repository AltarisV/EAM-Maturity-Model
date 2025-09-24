import streamlit as st
import pandas as pd
import altair as alt
import random
import re
from datetime import datetime
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

import matplotlib.pyplot as plt

st.set_page_config(page_title="EAM Maturity Assessment", layout="wide")

# ------------------------------
# Localization
# ------------------------------
translations = {
    "en": {
        "title": "EAM Maturity Assessment",
        "intro": """
This assessment is based on a maturity model for Enterprise Architecture Management (EAM).

For each dimension and phase of the ADM, criteria are shown that are assigned to a specific maturity level:

- If **all criteria** of a level and the levels below are met, this level is considered the **Baseline**.
- The highest level in which **at least one criterion** is met is considered the **Ceiling**.
- The actual maturity lies somewhere between the Baseline and the Ceiling.
- Within this range, the next steps to improve the Enterprise Architecture of the company should be planned (starting from the lowest level).

Please check all criteria that your organization currently meets.
""",
        "sidebar_tests": "Test functions",
        "btn_random": "🎲 Fill randomly",
        "btn_reset": "↩︎ Reset",
        "sidebar_chart": "Maturity Chart",
        "export": "Export",
        "docx_info": "`python-docx` is not installed. Please run: `pip install python-docx`.",
        "btn_docx": "📄 Create DOCX Report",
        "download_docx": "📥 Download DOCX",
        "results": "Assessment Results",
        "next_steps": "Next Steps",
        "no_next": "All criteria within the relevant range are fulfilled – no open Next Steps within the Baseline–Ceiling range.",
        "glossary": "ℹ️ Glossary / Explanations",
        "select_term": "Select a term",
        "lang_select": "🌐 Language",
        "chart-sidebar-heading": "Metric"
    },
    "de": {
        "title": "EAM Reifegrad-Assessment",
        "intro": """
Dieses Assessment basiert auf einem Reifegradmodell für Enterprise Architecture Management (EAM).

Für jede Dimension und Phase des ADM werden Kriterien angezeigt, die einem bestimmten Reifegrad-Level zugeordnet sind:

- Wenn **alle Kriterien** eines Levels und der darunterliegenden Levels erfüllt sind, gilt dieses Level als **Baseline**.
- Das höchste Level, in dem **mindestens ein Kriterium** erfüllt ist, gilt als **Deckel** (Ceiling).
- Die tatsächliche Reife liegt zwischen Baseline und Deckel.
- Innerhalb dieses Bereichs sollten die nächsten Schritte zur Verbesserung der Unternehmensarchitektur geplant werden (beginnend beim niedrigsten Level).

Bitte markieren Sie alle Kriterien, die Ihre Organisation aktuell erfüllt.
""",
        "sidebar_tests": "Testfunktionen",
        "btn_random": "🎲 Zufällig ausfüllen",
        "btn_reset": "↩︎ Zurücksetzen",
        "sidebar_chart": "Reifegrad-Diagramm",
        "export": "Export",
        "docx_info": "`python-docx` ist nicht installiert. Bitte ausführen: `pip install python-docx`.",
        "btn_docx": "📄 DOCX-Report erstellen",
        "download_docx": "📥 DOCX herunterladen",
        "results": "Bewertungsergebnisse",
        "next_steps": "Nächste Schritte",
        "no_next": "Alle Kriterien in den relevanten Bereichen sind erfüllt – keine offenen Next Steps im Baseline–Ceiling-Bereich.",
        "glossary": "ℹ️ Glossar / Erklärungen",
        "select_term": "Begriff auswählen",
        "lang_select": "🌐 Sprache",
        "chart-sidebar-heading": "Kennzahl"
    }
}

# ------------------------------
# Language state & toggle (early)
# ------------------------------
if "lang" not in st.session_state:
    st.session_state["lang"] = "en"

with st.sidebar:
    st.markdown("### 🌐 Language / Sprache")
    btn_label = "🇬🇧 English" if st.session_state["lang"] == "en" else "🇩🇪 Deutsch"
    if st.button(btn_label, key="lang_toggle"):
        # --- Antworten snappen (alle Checkbox-Werte sichern) ---
        snap = {}
        for k, v in st.session_state.items():
            if isinstance(k, str) and k.startswith("resp|"):
                snap[k] = bool(v)
        st.session_state["__resp_snapshot"] = snap

        # Sprache umschalten und rerun
        st.session_state["lang"] = "de" if st.session_state["lang"] == "en" else "en"
        st.rerun()

lang = st.session_state["lang"]
texts = translations[lang]


# ------------------------------
# Daten laden aus EINER Datei: alba.csv (ID + EN/DE-Text)
# ------------------------------
@st.cache_data(show_spinner=False)
def load_model(alba_path: str, lang: str) -> pd.DataFrame:
    """
    Erwartete Spalten in alba.csv:
      Dimension;ADM-Phases;Maturity Level;ID;Description_EN;Description_DE
    Gibt zurück:
      Dimension, ADM-Phases, level_num, ID (str), Description (sprachspezifisch mit Fallback)
    """
    df = pd.read_csv(alba_path, sep=";", encoding="utf-8-sig")

    required = {"Dimension", "ADM-Phases", "Maturity Level", "ID", "Description_EN", "Description_DE"}
    missing = required - set(df.columns)
    if missing:
        st.error(f"alba.csv fehlt Spalten: {', '.join(sorted(missing))}")
        st.stop()

    # Normalisieren
    df["Dimension"] = df["Dimension"].ffill()
    df["ADM-Phases"] = df.groupby("Dimension")["ADM-Phases"].ffill().fillna("")
    df["level_num"] = df["Maturity Level"].str.extract(r"(\d+)").astype("Int64")

    # IDs IMMER als String -> stabiler Streamlit-Key
    df["ID"] = df["ID"].astype(str)

    # Sprachtext mit Fallback (damit identische ID-Menge in beiden Sprachen sichtbar bleibt)
    def pick_desc(row):
        de = str(row.get("Description_DE", "") or "").strip()
        en = str(row.get("Description_EN", "") or "").strip()
        if lang == "de":
            return de if de else en  # DE bevorzugen, sonst EN
        else:
            return en if en else de  # EN bevorzugen, sonst DE

    df["Description"] = df.apply(pick_desc, axis=1)

    # Nur Level > 0 und nicht-leere Beschreibungen anzeigen
    df = df[(df["level_num"] > 0) & (df["Description"].astype(str).str.strip() != "")]
    return df[["Dimension", "ADM-Phases", "level_num", "ID", "Description"]]


def normalize_phase(p: str) -> str:
    """Phasen-String auf das kanonische Label bringen (kleine Toleranzen)."""
    if p is None:
        return ""
    s = str(p).strip()
    # häufiges Synonym abfangen: "D – ..." -> "B, C, D – ..."
    if s == "D – Business, Information Systems and Technology Architecture":
        return "B, C, D – Business, Information Systems and Technology Architecture"
    return s


def load_value_data(path: str, lang: str) -> tuple[dict, dict]:
    """
    Lädt Mehrwert-Infos.
    Unterstützt zwei Formen:
      1) ID-basiert:   Spalten: ID + (Value_EN/Value_DE oder Value)
      2) Tripel-basiert: Dimension; ADM-Phases; Maturity Level; (Value_EN/Value_DE oder Value)

    Rückgabe:
      (id_to_value: dict[str,str], triple_to_value: dict[tuple[str,str,int], str])
    """
    try:
        vdf = pd.read_csv(path, sep=";", encoding="utf-8-sig")
    except FileNotFoundError:
        return {}, {}

    # Spalten-Namen tolerant behandeln (lowercase-map)
    colmap = {c.lower(): c for c in vdf.columns}

    def has(name: str) -> bool:
        return name in colmap

    # passende Value-Spalte je Sprache finden
    candidates_de = ["value_de", "mehrwert", "mehrwert_de", "value"]
    candidates_en = ["value_en", "mehrwert_en", "value"]
    value_col_lc = None
    if lang == "de":
        for c in candidates_de:
            if has(c):
                value_col_lc = c
                break
    else:
        for c in candidates_en:
            if has(c):
                value_col_lc = c
                break
    if not value_col_lc:
        # letzte Chance: irgendeine Spalte mit "value" oder "mehrwert"
        for c in colmap:
            if "value" in c or "mehrwert" in c:
                value_col_lc = c
                break
    if not value_col_lc:
        return {}, {}

    value_col = colmap[value_col_lc]

    id_to_value: dict[str, str] = {}
    triple_to_value: dict[tuple[str, str, int], str] = {}

    # Fall 1: ID-basiert
    if has("id"):
        idcol = colmap["id"]
        tmp = vdf[[idcol, value_col]].copy()
        tmp[idcol] = tmp[idcol].astype(str).str.strip()
        tmp[value_col] = tmp[value_col].astype(str).fillna("").str.strip()
        id_to_value = dict(zip(tmp[idcol], tmp[value_col]))
        # zusätzlich versuchen wir auch Tripel zu bauen, falls vorhanden
        if has("dimension") and has("adm-phases") and has("maturity level"):
            dcol = colmap["dimension"]
            pcol = colmap["adm-phases"]
            lcol = colmap["maturity level"]
            tmp2 = vdf[[dcol, pcol, lcol, value_col]].copy()
            tmp2[dcol] = tmp2[dcol].astype(str).str.strip()
            tmp2[pcol] = tmp2[pcol].map(normalize_phase)
            tmp2["level_num"] = tmp2[lcol].astype(str).str.extract(r"(\d+)").astype(int)
            tmp2[value_col] = tmp2[value_col].astype(str).fillna("").str.strip()
            for _, r in tmp2.iterrows():
                triple_to_value[(r[dcol], r[pcol], int(r["level_num"]))] = r[value_col]
        return id_to_value, triple_to_value

    # Fall 2: Tripel-basiert (ohne ID)
    if has("dimension") and has("adm-phases") and has("maturity level"):
        dcol = colmap["dimension"]
        pcol = colmap["adm-phases"]
        lcol = colmap["maturity level"]
        tmp = vdf[[dcol, pcol, lcol, value_col]].copy()
        tmp[dcol] = tmp[dcol].astype(str).str.strip()
        tmp[pcol] = tmp[pcol].map(normalize_phase)
        tmp["level_num"] = tmp[lcol].astype(str).str.extract(r"(\d+)").astype(int)
        tmp[value_col] = tmp[value_col].astype(str).fillna("").str.strip()
        for _, r in tmp.iterrows():
            triple_to_value[(r[dcol], r[pcol], int(r["level_num"]))] = r[value_col]
        return {}, triple_to_value

    # nichts Passendes gefunden
    return {}, {}


try:
    raw = load_model("alba.csv", lang)
except Exception as e:
    st.error(f"Fehler beim Laden von alba.csv: {e}")
    st.stop()

# Gruppierte Kriterien je (Dimension, Phase, Level) – IDs & Descriptions als Listen
criteria = (
    raw.sort_values(["Dimension", "ADM-Phases", "level_num", "Description"])
        .groupby(["Dimension", "ADM-Phases", "level_num"])
        .agg(IDs=("ID", list), Descs=("Description", list))
        .reset_index()
)

# Reihenfolge der Phasen definieren (Labels aus dem Modell – hier englische Bezeichnungen)
phase_order = [
    "Preliminary",
    "A – Architecture Vision",
    "B, C, D – Business, Information Systems and Technology Architecture",
    "E – Opportunities & Solutions",
    "F – Migration Planning",
    "G – Implementation Governance",
    "H – Architecture Change Management",
    ""  # for Architecture Requirements Management without phase
]


def normalize_phase_name(phase: str) -> str:
    if phase is None:
        return ""
    s = re.sub(r"\s+", " ", str(phase)).strip()
    if not s or s in ["-", "—"]:
        return ""
    # Sammelphase B,C,D robust erkennen
    if "Business, Information Systems and Technology Architecture" in s:
        return "B, C, D – Business, Information Systems and Technology Architecture"
    # Requirements Mgmt hat keine Phase -> leer
    if s.lower().startswith("architecture requirements management"):
        return ""
    return s


st.title(texts["title"])
st.markdown(texts["intro"])

criteria["phase_order"] = criteria["ADM-Phases"].apply(
    lambda x: phase_order.index(x) if x in phase_order else len(phase_order)
)
criteria = criteria.sort_values(["phase_order", "level_num"]).reset_index(drop=True)

# Auswahl von vorheriger Sprache übernehmen (Snapshot zurückspielen)
if "__resp_snapshot" in st.session_state:
    resp_snap = st.session_state["__resp_snapshot"] or {}
    for _, row in criteria.iterrows():
        for item_id in row["IDs"]:
            k = f"resp|{item_id}"
            if k in resp_snap:
                st.session_state[k] = bool(resp_snap[k])
    del st.session_state["__resp_snapshot"]

id_to_value, triple_to_value = load_value_data("mehrwert.csv", lang)


def value_for_group(dim: str, phase: str, lvl: int, ids: list[str]) -> list[str]:
    # 1) Tripel bevorzugen (ein Wert für die ganze Gruppe)
    key = (str(dim).strip(), normalize_phase(phase), int(lvl))
    if key in triple_to_value:
        v = triple_to_value[key]
        return [v for _ in ids]
    # 2) Fallback pro ID
    out = []
    for i in ids:
        out.append(id_to_value.get(str(i).strip(), ""))
    return out


criteria["Values"] = criteria.apply(
    lambda r: value_for_group(r["Dimension"], r["ADM-Phases"], r["level_num"], r["IDs"]),
    axis=1
)

# ------------------------------
# Helpers for state, evaluation & export
# ------------------------------
RESP_KEY_PREFIX = "resp|"

LEVEL_FILL_PROB = {1: 0.90, 2: 0.80, 3: 0.50, 4: 0.10, 5: 0.02}


def fill_probability(level: int) -> float:
    return LEVEL_FILL_PROB.get(int(level), 0.50)


def checkbox_key(item_id: str) -> str:
    # Sprachunabhängig stabil
    return f"{RESP_KEY_PREFIX}{item_id}"


def init_state_if_missing():
    for _, row in criteria.iterrows():
        for item_id in row["IDs"]:
            k = checkbox_key(item_id)
            if k not in st.session_state:
                st.session_state[k] = False


def collect_responses() -> pd.DataFrame:
    records = []
    for _, row in criteria.iterrows():
        dim, phase, lvl = row["Dimension"], row["ADM-Phases"], row["level_num"]
        for item_id, desc in zip(row["IDs"], row["Descs"]):
            k = checkbox_key(item_id)
            records.append({
                "ID": item_id,
                "Dimension": dim,
                "ADM-Phases": phase,
                "level_num": lvl,
                "Description": desc,
                "Checked": bool(st.session_state.get(k, False)),
            })
    return pd.DataFrame(records)


def summarize(responses_df: pd.DataFrame):
    # Aggregate per (Dimension, Phase, Level)
    grp = (responses_df.groupby(["Dimension", "ADM-Phases", "level_num"])
           .agg(total=("Checked", "count"), done=("Checked", "sum"))
           .reset_index())
    grp["fulfilled"] = grp["done"] == grp["total"]
    grp["any"] = grp["done"] > 0

    # Results per (Dimension, Phase)
    results = []
    for (dim, phase), sub in grp.groupby(["Dimension", "ADM-Phases"], sort=False):
        baseline = 0
        for k in sorted(sub["level_num"].unique()):
            if sub.loc[sub["level_num"] <= k, "fulfilled"].all():
                baseline = k
        ceiling = sub.loc[sub["any"], "level_num"].max() if sub["any"].any() else 0
        results.append({
            "Dimension": dim,
            "ADM-Phases": phase,
            "Baseline": baseline,
            "Ceiling": ceiling,
        })

    df_res = pd.DataFrame(results)

    # Label (Phase, sonst Dimension)
    df_res["Label"] = df_res.apply(
        lambda r: r["ADM-Phases"] if r["ADM-Phases"] else r["Dimension"], axis=1
    )

    # Dynamische Sortierung: erst nach Phasenrang, dann Label
    def _phase_rank(phase):
        return phase_order.index(phase) if phase in phase_order else len(phase_order)

    df_res["__rank"] = df_res["ADM-Phases"].apply(_phase_rank)
    df_res = (
        df_res.sort_values(["__rank", "Label"])
            .drop(columns="__rank")
            .reset_index(drop=True)
    )

    # Zusatzkennzahl
    df_res["Average"] = (df_res["Baseline"] + df_res["Ceiling"]) / 2

    # Reihenfolge für die Altair-X-Achse
    x_order = df_res["Label"].tolist()
    return df_res, grp, x_order


def build_next_steps(df_res: pd.DataFrame, grp_levels: pd.DataFrame, responses_df: pd.DataFrame) -> pd.DataFrame:
    # Build Next Steps per phase: unmet criteria between Baseline+1 .. Ceiling (or Level 1 if nothing is met)
    next_rows = []
    for _, r in df_res.iterrows():
        dim, phase, baseline, ceiling = r["Dimension"], r["ADM-Phases"], r["Baseline"], r["Ceiling"]
        if ceiling == 0:
            target_levels = [1]
        else:
            target_levels = list(range(max(1, baseline + 1), ceiling + 1))

        for lvl in target_levels:
            crits = responses_df[(responses_df["Dimension"] == dim) &
                                 (responses_df["ADM-Phases"] == phase) &
                                 (responses_df["level_num"] == lvl)]
            if crits.empty:
                continue
            for _, row in crits.iterrows():
                if not row["Checked"]:
                    next_rows.append({
                        "Dimension": dim,
                        "ADM-Phases": phase if phase else "(no phase)",
                        "Level": int(lvl),
                        "ToDo": row["Description"],
                    })

    df_next = pd.DataFrame(next_rows).sort_values(["Dimension", "ADM-Phases", "Level"]).reset_index(drop=True)
    return df_next


def generate_chart_image(df_res: pd.DataFrame) -> BytesIO:
    """Render a compact chart with numeric indices on the x-axis for readability.
    The detailed mapping is provided in a table below the chart in the DOCX.
    """
    fig, ax = plt.subplots(figsize=(10, 4))

    x = list(range(1, len(df_res) + 1))
    baseline = df_res["Baseline"].tolist()
    ceiling = df_res["Ceiling"].tolist()

    ax.plot(x, baseline, marker="o", label="Baseline")
    ax.plot(x, ceiling, marker="o", label="Ceiling")

    ax.set_xticks(x)
    ax.set_xticklabels([str(i) for i in x])
    ax.set_ylabel("Level")
    ax.set_xlabel("Index (see table below)")
    ax.legend()
    ax.grid(True, axis="y", linestyle=":", alpha=0.4)

    buf = BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format='png', dpi=200)
    plt.close(fig)
    buf.seek(0)
    return buf


# ------------------------------
# Markdown → DOCX (bold + bullets + linebreaks)
# ------------------------------
def _add_runs_with_markdown(paragraph, text: str):
    """Add runs to a python-docx paragraph, interpreting **bold** segments."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part == "":
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_markdownish_text(doc, text: str):
    """Render a small subset of Markdown-like text into python-docx:
    - Empty lines → new paragraph
    - Lines starting with '- ' → bullet list paragraphs
    - '**bold**' → bold runs
    """
    for raw_line in text.split("\n"):
        line = raw_line.rstrip("\r")
        if line.strip() == "":
            doc.add_paragraph("")
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_markdown(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_runs_with_markdown(p, line)


def set_repeat_table_header(row):
    """Wiederholt die Kopfzeile einer Tabelle auf jeder Seite."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def prevent_row_split(row):
    """Verhindert Seitenumbruch innerhalb einer Tabellenzeile (no row split)."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def keep_with_next(paragraph, together=False):
    """Überschrift/Text nicht am Seitenende hängen lassen."""
    pf = paragraph.paragraph_format
    pf.keep_with_next = True
    if together:
        pf.keep_together = True


def build_docx_report(df_res: pd.DataFrame, responses_df: pd.DataFrame) -> BytesIO:
    if not DOCX_AVAILABLE:
        raise RuntimeError("`python-docx` is not installed.")

    # --- i18n nur für den DOCX-Export ---
    docx_t = {
        "en": {
            "title": "EAM Maturity Assessment",
            "generated_at": "Generated at:",
            "overview": "Maturity Overview",
            "idx_hint": "Indices on the chart correspond to the first column (#) in the table below.",
            "table_headers": ["#", "Phase / Dimension", "Baseline", "Ceiling", "Average"],
            "details": "Details & Next Steps",
            "baseline_lbl": "Baseline:",
            "ceiling_lbl": "Ceiling:",
            "reach_lead": "To reach Level {lvl}, the following criteria must be met:",
            "no_unmet": "All criteria for the immediate next level appear to be already met or no criteria are defined.",
            "intro_md": (
                "This assessment is based on the ALBA maturity model for Enterprise Architecture Management (EAM).\n"
                "For each dimension and phase of the ADM, criteria are shown that are assigned to a specific maturity level:\n"
                "- If **all criteria** of a level and the levels below are met, this level is considered the **Baseline**.\n"
                "- The highest level in which **at least one criterion** is met is considered the **Ceiling**.\n"
                "- The actual maturity lies somewhere between the Baseline and the Ceiling.\n"
                "- Within this range, the next steps to improve the Enterprise Architecture of the company should be planned (starting from the lowest level).\n"
            ),
        },
        "de": {
            "title": "EAM Reifegrad-Assessment",
            "generated_at": "Erstellt am:",
            "overview": "Reifegrad-Übersicht",
            "idx_hint": "Die Indizes in der Grafik entsprechen der ersten Spalte (#) in der untenstehenden Tabelle.",
            "table_headers": ["#", "Phase / Dimension", "Baseline", "Deckel", "Durchschnitt"],
            "details": "Details & Nächste Schritte",
            "baseline_lbl": "Baseline:",
            "ceiling_lbl": "Deckel:",
            "reach_lead": "Um Level {lvl} zu erreichen, müssen folgende Kriterien erfüllt sein:",
            "no_unmet": "Alle Kriterien für das unmittelbar nächste Level scheinen bereits erfüllt zu sein oder es sind keine Kriterien definiert.",
            "intro_md": (
                "Dieses Assessment basiert auf dem ALBA-Reifegradmodell für Enterprise Architecture Management (EAM).\n"
                "Für jede Dimension und ADM-Phase werden Kriterien gezeigt, die einem bestimmten Reifegrad zugeordnet sind:\n"
                "- Wenn **alle Kriterien** eines Levels und der darunterliegenden Levels erfüllt sind, gilt dieses Level als **Baseline**.\n"
                "- Das höchste Level, in dem **mindestens ein Kriterium** erfüllt ist, gilt als **Deckel**.\n"
                "- Die tatsächliche Reife liegt zwischen Baseline und Deckel.\n"
                "- Innerhalb dieses Bereichs sollten die nächsten Schritte zur Verbesserung der Unternehmensarchitektur geplant werden (beginnend beim niedrigsten Level).\n"
            ),
        },
    }["de" if lang == "de" else "en"]

    doc = Document()

    # --- Seitenlayout & Grundschrift ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    base = doc.styles["Normal"].font
    base.name = "Calibri"
    base.size = Pt(11)

    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(12)

    # --- Titel & Intro ---
    title = doc.add_heading(docx_t["title"], level=1)
    keep_with_next(title)
    stamp = doc.add_paragraph(f"{docx_t['generated_at']} {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    keep_with_next(stamp)

    add_markdownish_text(doc, docx_t["intro_md"])

    # --- Chart ---
    h2 = doc.add_heading(docx_t["overview"], level=2)
    keep_with_next(h2)

    chart_png = generate_chart_image(df_res)  # Legenden im Bild bleiben wie im Code (Baseline/Ceiling)
    pic_par = doc.add_paragraph()
    run = pic_par.add_run()
    run.add_picture(chart_png, width=Inches(6.5))
    pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_idx = doc.add_paragraph(docx_t["idx_hint"])
    keep_with_next(p_idx)

    # --- Tabelle ---
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, txt in enumerate(docx_t["table_headers"]):
        hdr[i].text = txt

    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for c in table.rows[0].cells:
        for p in c.paragraphs:
            if p.runs:
                p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, r in enumerate(df_res.itertuples(index=False), 1):
        row = table.add_row()
        prevent_row_split(row)
        row.cells[0].text = str(i)
        row.cells[1].text = str(getattr(r, "Label"))
        row.cells[2].text = str(int(getattr(r, "Baseline")))
        row.cells[3].text = str(int(getattr(r, "Ceiling")))
        row.cells[4].text = f"{float(getattr(r, 'Average')):.1f}"
        for j in [0, 2, 3, 4]:
            for p in row.cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    col_widths = [0.6, 3.9, 0.9, 0.9, 0.9]
    for row in table.rows:
        for j, w in enumerate(col_widths):
            row.cells[j].width = Inches(w)

    # --- Nächster Abschnitt ---
    h2 = doc.add_heading(docx_t["details"], level=2)
    keep_with_next(h2)

    for _, r in df_res.iterrows():
        label = str(r["Label"])
        dim = r["Dimension"]
        phase = r["ADM-Phases"]
        baseline = int(r["Baseline"])
        ceiling = int(r["Ceiling"])

        h3 = doc.add_heading(label, level=3)
        keep_with_next(h3)

        p = doc.add_paragraph()
        p.add_run(f"{docx_t['baseline_lbl']} ").bold = True
        p.add_run(str(baseline))
        p.add_run(f"; {docx_t['ceiling_lbl']} ").bold = True
        p.add_run(str(ceiling))
        keep_with_next(p)

        target_level = max(1, baseline + 1)
        if 0 < ceiling < target_level:
            target_level = baseline + 1

        lead = doc.add_paragraph(docx_t["reach_lead"].format(lvl=target_level))
        keep_with_next(lead)

        crits = responses_df[
            (responses_df["Dimension"] == dim) &
            (responses_df["ADM-Phases"] == phase) &
            (responses_df["level_num"] == target_level)
            ]
        unmet = crits[~crits["Checked"]]

        if unmet.empty:
            doc.add_paragraph(docx_t["no_unmet"])
        else:
            for _, row in unmet.iterrows():
                li = doc.add_paragraph(row["Description"], style="List Bullet")
                keep_with_next(li, together=True)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ------------------------------
# Sidebar: Test functions, chart & export
# ------------------------------
with st.sidebar:
    st.subheader(texts["sidebar_tests"])
    col_a, col_b = st.columns(2)

    # Zufällig ausfüllen (mit stabilen ID-Keys)
    with col_a:
        if st.button(texts["btn_random"]):
            init_state_if_missing()
            for _, row in criteria.iterrows():
                lvl = int(row["level_num"])
                p = fill_probability(lvl)
                for item_id in row["IDs"]:
                    st.session_state[checkbox_key(item_id)] = (random.random() < p)
            st.rerun()

    # Antworten zurücksetzen (Sprache bleibt unangetastet)
    with col_b:
        if st.button(texts["btn_reset"]):
            for k in list(st.session_state.keys()):
                if isinstance(k, str) and k.startswith(RESP_KEY_PREFIX):
                    st.session_state[k] = False
            st.rerun()

# ------------------------------
# UI: Render checklists
# ------------------------------
init_state_if_missing()
for _, row in criteria.iterrows():
    dim, phase, level = row["Dimension"], row["ADM-Phases"], row["level_num"]
    header = f"{dim} – {phase} – Level {level}" if phase else f"{dim} – Level {level}"

    with st.expander(header, expanded=False):
        first_id = row["IDs"][0] if row["IDs"] else None
        for item_id, desc, val in zip(row["IDs"], row["Descs"], row["Values"]):
            desc_str = str(desc).strip()
            if not desc_str:
                continue
            k = checkbox_key(item_id)
            if item_id == first_id:
                col1, col2 = st.columns([20, 1])
                with col1:
                    st.checkbox(desc_str, key=k)
                with col2:
                    with st.popover("ℹ️"):
                        s = str(val).strip()
                        if not s:
                            s = "Value not available." if lang == "en" else "Mehrwert nicht verfügbar."
                        st.markdown(s)
            else:
                st.checkbox(desc_str, key=k)

# ------------------------------
# Evaluation & visualization
# ------------------------------
responses_df = collect_responses()
df_res, grp_levels, x_order = summarize(responses_df)

# Chart in sidebar (Altair)
chart = alt.Chart(df_res).transform_fold(
    fold=["Baseline", "Ceiling"],
    as_=["Metric", "Level"]
).mark_line(point=True).encode(
    x=alt.X("Label:N", title="Phase / Dimension", sort=x_order),
    y=alt.Y("Level:Q", title="Level"),
    color=alt.Color("Metric:N", title=texts["chart-sidebar-heading"]),
    tooltip=[
        alt.Tooltip("Label:N", title="Phase/Dimension"),
        alt.Tooltip("Metric:N", title=texts["chart-sidebar-heading"]),
        alt.Tooltip("Level:Q", title="Level")
    ]
)

with st.sidebar:
    st.subheader(texts["sidebar_chart"])
    st.altair_chart(chart, use_container_width=True)

    st.markdown("---")
    st.subheader(texts["export"])
    if not DOCX_AVAILABLE:
        st.info(texts["docx_info"])
    else:
        if st.button(texts["btn_docx"]):
            try:
                docx_buf = build_docx_report(df_res, responses_df)
                st.download_button(
                    label=texts["download_docx"],
                    data=docx_buf.getvalue(),
                    file_name=f"eam_maturity_report_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"Error while creating DOCX: {e}")

# Hauptbereich / Main area
st.subheader(texts["results"])
st.dataframe(df_res, use_container_width=True)

st.subheader(texts["next_steps"])
df_next = build_next_steps(df_res, grp_levels, responses_df)
if df_next.empty:
    st.success(texts["no_next"])
else:
    st.dataframe(df_next, use_container_width=True)

# Glossar zuletzt (optional im Hauptbereich)
with st.expander(texts["glossary"]):
    glossary = {
        "Baseline": "Highest level where all criteria up to and including that level are fulfilled.",
        "Ceiling": "Highest level where at least one criterion is fulfilled.",
        "EAM": "Enterprise Architecture Management — holistic planning and governance of the enterprise architecture.",
        "ADM": "Architecture Development Method — the TOGAF method with phases from Preliminary to H.",
        "Architecture Requirements Management": "Cross-cutting process that manages requirements across all phases.",
    }
    term = st.selectbox(texts["select_term"],
                        options=["(bitte wählen)" if lang == "de" else "(please choose)"] + list(glossary.keys()))
    if term not in ["(bitte wählen)", "(please choose)"]:
        st.markdown(f"**{term}:** {glossary[term]}")
