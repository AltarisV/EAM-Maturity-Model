"""
exports.py
Chart rendering plus DOCX and Excel exports.
"""

from io import BytesIO
from datetime import datetime
import re
import pandas as pd
import matplotlib.pyplot as plt

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# ---------- Charts ----------

def generate_chart_image(df_res: pd.DataFrame) -> BytesIO:
    """
    Render a compact Baseline/Ceiling line chart and return it as PNG bytes.

    Args:
        df_res: Summary table from `core.summarize`.

    Returns:
        BytesIO positioned at start, containing a PNG image.
    """
    fig, ax = plt.subplots(figsize=(10, 4))
    x = list(range(1, len(df_res) + 1))
    ax.plot(x, df_res["Baseline"].tolist(), marker="o", label="Baseline")
    ax.plot(x, df_res["Ceiling"].tolist(), marker="o", label="Ceiling")
    ax.set_xticks(x)
    ax.set_xticklabels([str(i) for i in x])
    ax.set_ylabel("Level")
    ax.set_xlabel("Index (see table below)")
    ax.legend()
    ax.grid(True, axis="y", linestyle=":", alpha=0.4)
    buf = BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format='png', dpi=200)
    plt.close(fig)
    buf.seek(0)
    return buf


# ---------- Minimal Markdown support for DOCX ----------

def _add_runs_with_markdown(paragraph, text: str) -> None:
    """
    Add text to a python-docx paragraph honoring **bold** segments.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_markdownish_text(doc, text: str) -> None:
    """
    Render a small subset of Markdown:
      - Empty line → new paragraph
      - Lines starting with '- ' → bullet list
      - '**bold**' → bold text
    """
    for raw_line in text.split("\n"):
        line = raw_line.rstrip("\r")
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_markdown(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_runs_with_markdown(p, line)


def set_repeat_table_header(row) -> None:
    """Mark a table header row to repeat on each page in DOCX."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def prevent_row_split(row) -> None:
    """Prevent Word from splitting a table row across pages."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def keep_with_next(paragraph, together: bool = False) -> None:
    """Keep a paragraph with the following paragraph; optionally keep together."""
    pf = paragraph.paragraph_format
    pf.keep_with_next = True
    if together:
        pf.keep_together = True


# ---------- DOCX export ----------

def build_docx_report(df_res: pd.DataFrame, responses_df: pd.DataFrame, lang: str) -> BytesIO:
    """
    Build a DOCX report: intro, chart, summary table, and per-phase next steps.

    Args:
        df_res: Summary table from `core.summarize`.
        responses_df: Row-per-criterion table from `core.collect_responses`.
        lang: "en" or "de" – controls labels inside the document.

    Returns:
        BytesIO positioned at start, containing a .docx file.

    Raises:
        RuntimeError: if python-docx is not available.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("`python-docx` is not installed.")

    t = {
        "en": {
            "title": "EAM Maturity Assessment",
            "generated_at": "Generated at:",
            "overview": "Maturity Overview",
            "idx_hint": "Indices on the chart correspond to the first column (#) in the table below.",
            "table_headers": ["#", "Phase / Dimension", "Baseline", "Ceiling", "Average"],
            "details": "Details & Next Steps",
            "baseline_lbl": "Baseline:",
            "ceiling_lbl": "Ceiling:",
            "reach_lead": "To reach Level {lvl}, the following criteria must be met:",
            "no_unmet": "All criteria for the immediate next level appear to be already met or no criteria are defined.",
            "intro_md": (
                "This assessment is based on the ALBA maturity model for Enterprise Architecture Management (EAM).\n"
                "- If **all criteria** of a level and the levels below are met, this level is considered the "
                "**Baseline**.\n "
                "- The highest level in which **at least one criterion** is met is considered the **Ceiling**.\n"
                "- Within this range, the next steps should be planned (starting from the lowest level).\n"
            ),
        },
        "de": {
            "title": "EAM Reifegrad-Assessment",
            "generated_at": "Erstellt am:",
            "overview": "Reifegrad-Übersicht",
            "idx_hint": "Die Indizes in der Grafik entsprechen der ersten Spalte (#) in der untenstehenden Tabelle.",
            "table_headers": ["#", "Phase / Dimension", "Baseline", "Deckel", "Durchschnitt"],
            "details": "Details & Nächste Schritte",
            "baseline_lbl": "Baseline:",
            "ceiling_lbl": "Deckel:",
            "reach_lead": "Um Level {lvl} zu erreichen, müssen folgende Kriterien erfüllt sein:",
            "no_unmet": "Alle Kriterien für das unmittelbar nächste Level scheinen bereits erfüllt zu sein oder es "
                        "sind keine Kriterien definiert.",
            "intro_md": (
                "Dieses Assessment basiert auf dem ALBA-Reifegradmodell für Enterprise Architecture Management (EAM).\n"
                "- Wenn **alle Kriterien** eines Levels und der darunterliegenden Levels erfüllt sind, gilt dieses "
                "Level als **Baseline**.\n "
                "- Das höchste Level, in dem **mindestens ein Kriterium** erfüllt ist, gilt als **Deckel**.\n"
                "- Innerhalb dieses Bereichs sollten die nächsten Schritte geplant werden (beginnend beim niedrigsten "
                "Level).\n "
            ),
        },
    }["de" if lang == "de" else "en"]

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    base = doc.styles["Normal"].font
    base.name = "Calibri"
    base.size = Pt(11)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(12)

    title = doc.add_heading(t["title"], level=1)
    keep_with_next(title)
    stamp = doc.add_paragraph(f"{t['generated_at']} {datetime.now():%Y-%m-%d %H:%M:%S}")
    keep_with_next(stamp)
    add_markdownish_text(doc, t["intro_md"])

    h2 = doc.add_heading(t["overview"], level=2)
    keep_with_next(h2)
    chart_png = generate_chart_image(df_res)
    pic_par = doc.add_paragraph()
    run = pic_par.add_run()
    run.add_picture(chart_png, width=Inches(6.5))
    pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_idx = doc.add_paragraph(t["idx_hint"])
    keep_with_next(p_idx)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, txt in enumerate(t["table_headers"]):
        hdr[i].text = txt
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for c in table.rows[0].cells:
        for p in c.paragraphs:
            if p.runs:
                p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, r in enumerate(df_res.itertuples(index=False), 1):
        row = table.add_row()
        prevent_row_split(row)
        row.cells[0].text = str(i)
        row.cells[1].text = str(getattr(r, "Label"))
        row.cells[2].text = str(int(getattr(r, "Baseline")))
        row.cells[3].text = str(int(getattr(r, "Ceiling")))
        row.cells[4].text = f"{float(getattr(r, 'Average')):.1f}"
        for j in [0, 2, 3, 4]:
            for p in row.cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    col_widths = [0.6, 3.9, 0.9, 0.9, 0.9]
    for row in table.rows:
        for j, w in enumerate(col_widths):
            row.cells[j].width = Inches(w)

    h2 = doc.add_heading(t["details"], level=2)
    keep_with_next(h2)

    for _, r in df_res.iterrows():
        label = str(r["Label"])
        baseline = int(r["Baseline"])
        ceiling = int(r["Ceiling"])

        h3 = doc.add_heading(label, level=3)
        keep_with_next(h3)

        p = doc.add_paragraph()
        p.add_run(f"{t['baseline_lbl']} ").bold = True
        p.add_run(str(baseline))
        p.add_run(f"; {t['ceiling_lbl']} ").bold = True
        p.add_run(str(ceiling))
        keep_with_next(p)

        target_level = 1 if ceiling == 0 else max(1, baseline + 1)

        dim, phase = r["Dimension"], r["ADM-Phases"]
        crits = responses_df[
            (responses_df["Dimension"] == dim) &
            (responses_df["ADM-Phases"] == phase) &
            (responses_df["level_num"] == target_level)
            ]
        unmet = crits[~crits["Checked"]]
        if unmet.empty:
            doc.add_paragraph(t["no_unmet"])
        else:
            for _, row in unmet.iterrows():
                li = doc.add_paragraph(row["Description"], style="List Bullet")
                keep_with_next(li, together=True)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ---------- Excel export ----------

def _autosize_excel_columns(writer: pd.ExcelWriter, df: pd.DataFrame, sheet_name: str,
                            min_w: int = 8, max_w: int = 80, wrap_cols: list[str] | None = None) -> None:
    """Set Excel column widths from content length and optionally enable wrapping.

    Args:
        writer: Active Excel writer.
        df: Data to inspect for width.
        sheet_name: Target sheet name.
        min_w: Minimum column width.
        max_w: Maximum column width.
        wrap_cols: Columns that should use text-wrap formatting.
    """
    ws = writer.sheets[sheet_name]
    book = writer.book
    wrap_fmt = book.add_format({"text_wrap": True})
    for col_idx, col in enumerate(df.columns):
        series = df[col].astype(str)
        max_len = max([len(str(col))] + series.map(len).tolist())
        width = max(min_w, min(max_w, max_len + 2))
        if wrap_cols and col in wrap_cols:
            ws.set_column(col_idx, col_idx, width, wrap_fmt)
        else:
            ws.set_column(col_idx, col_idx, width)


def generate_excel_report(df_res: pd.DataFrame, responses_df: pd.DataFrame, lang: str) -> BytesIO:
    """
    Build an Excel workbook with:
        - Summary sheet,
        - Next steps sheet,
        - Raw responses,
        - Embedded chart image.

    Args:
        df_res: Summary table from `core.summarize`.
        responses_df: Row-per-criterion from `core.collect_responses`.
        lang: "en" or "de" – controls column labels and sheet names.

    Returns:
        BytesIO positioned at start, containing a .xlsx file.
    """
    from core import summarize, build_next_steps
    _, grp_levels, _ = summarize(responses_df)
    df_next = build_next_steps(df_res, grp_levels, responses_df)

    xl = {
        "en": {
            "sheet_summary": "Summary", "sheet_next": "NextSteps", "sheet_resp": "Responses", "sheet_chart": "Chart",
            "col_phase": "Phase", "col_label": "Phase / Dimension", "col_baseline": "Baseline",
            "col_ceiling": "Ceiling", "col_avg": "Average", "col_level": "Level",
            "col_desc": "Description", "col_checked": "Checked", "arm_label": "Architecture Requirements Management",
        },
        "de": {
            "sheet_summary": "Übersicht", "sheet_next": "Nächste Schritte", "sheet_resp": "Antworten", "sheet_chart": "Diagramm",
            "col_phase": "Phase", "col_label": "Phase / Dimension", "col_baseline": "Baseline",
            "col_ceiling": "Deckel", "col_avg": "Durchschnitt", "col_level": "Level",
            "col_desc": "Beschreibung", "col_checked": "Erfüllt", "arm_label": "Architecture Requirements Management",
        },
    }["de" if lang == "de" else "en"]

    df_summary = df_res.copy()
    if "ADM-Phases" in df_summary.columns:
        df_summary = df_summary.rename(columns={"ADM-Phases": xl["col_phase"]})
    if "Label" in df_summary.columns:
        df_summary = df_summary.rename(columns={"Label": xl["col_label"]})
    df_summary = df_summary.rename(columns={
        "Baseline": xl["col_baseline"], "Ceiling": xl["col_ceiling"], "Average": xl["col_avg"],
    })

    df_next_out = df_next.copy()
    if "ADM-Phases" in df_next_out.columns:
        df_next_out = df_next_out.rename(columns={"ADM-Phases": xl["col_phase"]})
    if xl["col_phase"] in df_next_out.columns:
        df_next_out[xl["col_phase"]] = df_next_out[xl["col_phase"]].replace("", xl["arm_label"])

    df_resp_out = responses_df.copy().rename(columns={
        "ADM-Phases": xl["col_phase"], "level_num": xl["col_level"], "Description": xl["col_desc"],
        "Checked": xl["col_checked"],
    })

    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        df_summary.to_excel(writer, sheet_name=xl["sheet_summary"], index=False)
        _autosize_excel_columns(writer, df_summary, xl["sheet_summary"], wrap_cols=[xl["col_label"]])

        df_next_out.to_excel(writer, sheet_name=xl["sheet_next"], index=False)
        _autosize_excel_columns(writer, df_next_out, xl["sheet_next"], wrap_cols=["ToDo", xl["col_phase"]])

        df_resp_out.to_excel(writer, sheet_name=xl["sheet_resp"], index=False)
        _autosize_excel_columns(writer, df_resp_out, xl["sheet_resp"], wrap_cols=[xl["col_desc"], xl["col_phase"]])

        chart_png = generate_chart_image(df_res)
        ws = writer.book.add_worksheet(xl["sheet_chart"])
        writer.sheets[xl["sheet_chart"]] = ws
        ws.insert_image(0, 0, "chart.png", {"image_data": chart_png})

    buf.seek(0)
    return buf
